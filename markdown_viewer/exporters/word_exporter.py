"""
Word document exporter using python-docx with Playwright-based rendering.

This exporter renders math formulas and Mermaid diagrams in a browser using
Playwright, captures them as images, and embeds them in the Word document.
This ensures the Word export looks exactly like the browser version.
"""

# pylint: disable=broad-exception-caught,too-many-lines

import logging
import re
import os
import base64
import tempfile
import hashlib
import urllib.request
import urllib.parse
from datetime import datetime
from typing import Optional, Dict, Any
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString
from playwright.sync_api import sync_playwright

logger = logging.getLogger(__name__)

# Constants
RENDER_WAIT_MS = 2500  # Wait for KaTeX and Mermaid to render


class WordExporter:
    """Export content to Word document with browser-quality rendering."""

    def __init__(self):
        """Initialize Word exporter."""
        self.playwright = None
        self.browser = None
        self.page = None
        self.temp_html_path = None
        self.screenshot_cache = {}  # Cache element screenshots by content hash
        self.math_counter = 0  # Track which math element we're processing
        self.mermaid_counter = 0  # Track which mermaid element we're processing

    def _ensure_browser(self):
        """Initialize Playwright browser if not already running."""
        if self.playwright is None:
            self.playwright = sync_playwright().start()
            self.browser = self.playwright.chromium.launch(headless=True)
            logger.info("Playwright browser launched for Word export")

    def _load_html(self, html_content: str):
        """Load HTML into browser for rendering."""
        self._ensure_browser()

        # Preprocess HTML to add unique IDs to math and mermaid elements
        soup = BeautifulSoup(html_content, "html.parser")

        # Add IDs to all math elements
        math_elements = soup.find_all(class_=re.compile(r"arithmatex|katex"))
        for idx, elem in enumerate(math_elements):
            elem["id"] = f"math_element_{idx}"

        # Add IDs to all mermaid diagrams
        mermaid_elements = soup.find_all(class_="mermaid")
        for idx, elem in enumerate(mermaid_elements):
            elem["id"] = f"mermaid_element_{idx}"

        # Convert local file images to data URLs BEFORE Playwright loads them
        # This is critical because browsers block loading local file:// paths for security
        img_elements = soup.find_all("img")
        for img in img_elements:
            src = img.get("src", "")
            if not src:
                continue

            # Skip remote URLs (http/https)
            if src.startswith(("http://", "https://")):
                continue

            # Skip if already a data URL
            if src.startswith("data:"):
                continue

            # It's a local file path - convert to data URL
            try:
                # Handle Windows paths with backslashes and URL encoding
                import urllib.parse
                from pathlib import Path

                # Decode URL encoding
                decoded_path = urllib.parse.unquote(src)
                logger.info("Processing image src: %s", src)
                logger.info("After URL decode: %s", decoded_path)

                # Convert to Path object
                img_path = Path(decoded_path)
                logger.info(
                    "Path object created: %s (absolute=%s)", img_path, img_path.is_absolute()
                )

                # If not absolute, resolve relative to markdown file directory
                if not img_path.is_absolute() and self.md_file_path:
                    img_path = Path(self.md_file_path).parent / decoded_path
                    logger.info("Resolved relative path to: %s", img_path)

                # Read and convert to base64
                logger.info(
                    "Checking if exists: %s, is_file: %s",
                    img_path.exists(),
                    img_path.is_file() if img_path.exists() else "N/A",
                )
                if img_path.exists() and img_path.is_file():
                    import base64

                    with open(img_path, "rb") as f:
                        img_data = f.read()

                    logger.info("Read %d bytes from image file", len(img_data))

                    # Detect MIME type
                    ext = img_path.suffix.lower()
                    mime_types = {
                        ".png": "image/png",
                        ".jpg": "image/jpeg",
                        ".jpeg": "image/jpeg",
                        ".gif": "image/gif",
                        ".webp": "image/webp",
                        ".svg": "image/svg+xml",
                        ".bmp": "image/bmp",
                    }
                    mime_type = mime_types.get(ext, "image/png")

                    # Convert to data URL
                    b64_data = base64.b64encode(img_data).decode("ascii")
                    data_url = f"data:{mime_type};base64,{b64_data}"

                    # Replace src
                    img["src"] = data_url
                    logger.info(
                        "✅ Converted local image to data URL: %s (%d KB)",
                        img_path.name,
                        len(img_data) // 1024,
                    )
                else:
                    logger.warning("❌ Image file not found or not accessible: %s", img_path)
            except Exception as e:
                logger.error(
                    "❌ Failed to convert image to data URL: %s - %s", src, e, exc_info=True
                )

        # Convert back to HTML
        html_with_ids = str(soup)

        # Create complete HTML document with scripts for KaTeX and Mermaid rendering
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.45/dist/katex.min.css">
    <script src="https://cdn.jsdelivr.net/npm/katex@0.16.45/dist/katex.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.min.js"></script>
    <style>
        body {{ font-family: sans-serif; padding: 20px; }}
        .mermaid {{ margin: 20px 0; padding: 10px; }}
        /* No padding on math elements for tight screenshots */
        .arithmatex {{ display: inline-block; }}
        .katex {{ }}
    </style>
</head>
<body>
{html_with_ids}
<script>
    // Initialize Mermaid
    mermaid.initialize({{
        startOnLoad: false,
        theme: 'default',
        securityLevel: 'loose'
    }});
    
    // Render KaTeX math
    document.querySelectorAll('.arithmatex').forEach(function(el) {{
        var math = el.textContent;
        // Remove LaTeX delimiters
        math = math.replace(/^\\\\\\[/, '').replace(/\\\\\\]$/, '');
        math = math.replace(/^\\\\\\(/, '').replace(/\\\\\\)$/, '');
        try {{
            katex.render(math.trim(), el, {{
                throwOnError: false,
                displayMode: el.tagName === 'DIV'
            }});
        }} catch(e) {{
            console.error('KaTeX error:', e);
        }}
    }});
    
    // Render Mermaid diagrams
    (async function() {{
        var diagrams = document.querySelectorAll('.mermaid');
        for (var i = 0; i < diagrams.length; i++) {{
            try {{
                var source = diagrams[i].textContent;
                var id = 'mermaid-' + Date.now() + '-' + i;
                var result = await mermaid.render(id, source);
                diagrams[i].innerHTML = result.svg;
            }} catch(e) {{
                console.error('Mermaid error:', e);
            }}
        }}
    }})();
</script>
</body>
</html>"""

        # Save HTML to temp file
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8"
        ) as f:
            f.write(full_html)
            self.temp_html_path = f.name

        # Create new page with high DPI for crisp screenshots
        context = self.browser.new_context(device_scale_factor=2)
        self.page = context.new_page()

        # Load HTML and wait for network to be idle (scripts loaded)
        self.page.goto(f"file://{self.temp_html_path}", wait_until="networkidle")

        # Wait for rendering (KaTeX, Mermaid async rendering)
        # Mermaid uses async await, so needs more time
        self.page.wait_for_timeout(5000)

        # Wait for KaTeX elements to appear (if any math exists)
        try:
            self.page.wait_for_selector(".katex, .mermaid svg", timeout=3000)
        except Exception:
            pass  # No math or diagrams present

        logger.info("HTML loaded and rendered in browser")

    def _screenshot_element(self, selector: str, content_hash: str) -> Optional[bytes]:
        """
        Take screenshot of an element by CSS selector.

        Args:
            selector: CSS selector for the element
            content_hash: Hash of element content for caching

        Returns:
            PNG image bytes or None if element not found
        """
        if content_hash in self.screenshot_cache:
            return self.screenshot_cache[content_hash]

        try:
            element = self.page.query_selector(selector)
            if element:
                # Scroll element into view to ensure it's fully visible
                element.scroll_into_view_if_needed()

                # Wait for any layout shifts
                self.page.wait_for_timeout(200)

                # Take high-quality screenshot with padding to avoid clipping
                # The padding parameter adds space around the element
                screenshot = element.screenshot(type="png", omit_background=True)
                self.screenshot_cache[content_hash] = screenshot
                logger.debug(f"Screenshot captured for {selector}: {len(screenshot)} bytes")
                return screenshot

            logger.warning(f"Element not found for selector: {selector}")
        except Exception as e:
            logger.warning("Failed to screenshot element %s: %s", selector, e)

        return None

    def _cleanup(self):
        """Clean up browser and temp files."""
        if self.page:
            try:
                self.page.close()
            except Exception as e:
                logger.warning("Error closing page: %s", e)
            self.page = None

        if self.browser:
            try:
                self.browser.close()
            except Exception as e:
                logger.warning("Error closing browser: %s", e)
            self.browser = None

        if self.playwright:
            try:
                self.playwright.stop()
            except Exception as e:
                logger.warning("Error stopping playwright: %s", e)
            self.playwright = None

        if self.temp_html_path and os.path.exists(self.temp_html_path):
            try:
                os.unlink(self.temp_html_path)
            except OSError as e:
                logger.warning("Failed to delete temp HTML: %s", e)

    def export(
        self,
        html_content: str,
        markdown_content: str,  # pylint: disable=unused-argument
        output_path: str,
        md_file_path: Optional[str] = None,
        options: Optional[Dict[str, Any]] = None,
    ) -> None:
        """
        Export content to Word document with pixel-perfect math and diagram rendering.

        Args:
            html_content: Full HTML string to render
            markdown_content: Original markdown (not used)
            output_path: Path to save .docx file
            md_file_path: Optional path to source markdown file for resolving relative images
            options: Optional export options
        """
        if options is None:
            options = {}

        # Store md_file_path for image resolution
        self.md_file_path = md_file_path

        try:
            # Load HTML into browser for screenshot capability
            self._load_html(html_content)

            # Get the rendered HTML from the browser (after KaTeX/Mermaid rendering)
            # This ensures we process the actual rendered elements, not the raw LaTeX
            rendered_html = self.page.content()

            # DEBUG: Check if images exist in HTML
            img_count = rendered_html.count("<img")
            logger.info("Rendered HTML contains %d <img tags", img_count)
            if img_count > 0:
                # Log first few img tags
                import re

                img_tags = re.findall(r"<img[^>]+>", rendered_html)
                for idx, tag in enumerate(img_tags[:3]):
                    logger.info("IMG tag %d: %s", idx, tag[:200])

            # Create Word document
            doc = Document()

            # Parse the RENDERED HTML (not the original)
            soup = BeautifulSoup(rendered_html, "html.parser")

            # Add TOC with clickable bookmarks
            self._add_toc(doc, soup)

            # Remove HTML TOC div to avoid duplication
            toc_div = soup.find("div", class_="toc")
            if toc_div:
                toc_div.decompose()

            # Process content
            self._process_element(doc, soup, is_root=True)

            # Add footer
            self._add_footer(doc)

            # Save document
            doc.save(output_path)
            logger.info("Word document saved: %s", output_path)

        finally:
            # Always clean up browser resources
            self._cleanup()

    def _add_toc(self, doc, soup):
        """Add Table of Contents with clickable bookmarks."""
        headings = soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
        if not headings:
            return

        # TOC Title
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_before = Pt(0)
        toc_para.paragraph_format.space_after = Pt(8)
        toc_run = toc_para.add_run("Table of Contents")
        toc_run.bold = True
        toc_run.font.size = Pt(16)
        toc_run.font.color.rgb = RGBColor(47, 84, 150)

        # TOC Entries
        for idx, heading in enumerate(headings):
            level = int(heading.name[1])
            text = heading.get_text().strip()

            if text:
                # Generate bookmark name
                bookmark_name = f"_Heading_{idx}"

                # Add TOC item with indentation
                toc_item = doc.add_paragraph()
                toc_item.paragraph_format.left_indent = Inches((level - 1) * 0.5)
                # Set tight spacing for TOC
                toc_item.paragraph_format.space_before = Pt(0)
                toc_item.paragraph_format.space_after = Pt(2)
                toc_item.paragraph_format.line_spacing = 1.0

                # Add hyperlink to bookmark
                self._add_internal_hyperlink(toc_item, bookmark_name, text)

                # Make h1/h2 bold
                if level <= 2:
                    for run in toc_item.runs:
                        run.bold = True

                # Add bookmark to the heading itself (will be done in _add_heading)
                heading.attrs["bookmark_name"] = bookmark_name

        # Spacing after TOC
        doc.add_paragraph()

    def _add_internal_hyperlink(self, paragraph, bookmark_name, text):
        """Add clickable internal hyperlink (to bookmark)."""
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Blue color and underline
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)  # pylint: disable=protected-access

    def _add_bookmark(self, paragraph, bookmark_name):
        """Add bookmark to a paragraph for TOC linking."""
        # Generate unique bookmark ID
        bookmark_id = str(hash(bookmark_name) % 10000)

        # Bookmark start
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), bookmark_id)
        bookmark_start.set(qn("w:name"), bookmark_name)

        # Bookmark end
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), bookmark_id)

        # Insert bookmark around paragraph
        p_element = paragraph._p  # pylint: disable=protected-access
        p_element.insert(0, bookmark_start)
        p_element.append(bookmark_end)

    def _add_footer(self, doc):
        """Add footer with generation timestamp."""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        current_datetime = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        footer_text = f"Generated by Markdown Viewer on {current_datetime}"

        run = footer_para.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _process_element(self, doc, element, is_root=False):
        """
        Recursively process HTML elements and add to Word document.

        Special handling for math (.arithmatex) and diagrams (.mermaid):
        - Screenshots are captured from rendered browser
        - Inserted as images for pixel-perfect accuracy
        """
        if is_root:
            for child in element.children:
                if hasattr(child, "name"):
                    self._process_element(doc, child)
            return

        # Math formulas (with screenshot)
        if element.name in ["span", "div"] and self._has_math_class(element):
            self._add_math_as_image(doc, element)
        # Mermaid diagrams (with screenshot)
        elif element.name == "div" and "mermaid" in element.get("class", []):
            self._add_mermaid_as_image(doc, element)
        # Headings
        elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            self._add_heading(doc, element, level)
        # Paragraphs
        elif element.name == "p":
            # Skip TOC placeholder
            if "toc-placeholder" in element.get("class", []):
                p = doc.add_paragraph(element.get_text(), style="Intense Quote")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.0
            else:
                self._add_paragraph_with_formatting(doc, element)
        # Lists
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                self._add_list_item(doc, li, style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                self._add_list_item(doc, li, style="List Number")
        # Code blocks
        elif element.name == "pre":
            self._add_code_block(doc, element)
        # Blockquotes
        elif element.name == "blockquote":
            self._add_paragraph_with_formatting(doc, element, style="Quote")
        # Tables
        elif element.name == "table":
            self._process_table(doc, element)
        # Images
        elif element.name == "img":
            self._add_image(doc, element)
        # Horizontal rules
        elif element.name == "hr":
            self._add_horizontal_rule(doc)
        # Containers - process children
        else:
            if hasattr(element, "children"):
                for child in element.children:
                    if hasattr(child, "name") and child.name:
                        self._process_element(doc, child)

    def _add_heading(self, doc, element, level):
        """Add heading with bookmark for TOC linking."""
        text = self._get_text_with_emojis(element)
        if not text.strip():
            return

        para = doc.add_heading(text, level=level)
        # Set tight spacing for headings
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

        # Add bookmark if set by TOC
        bookmark_name = element.attrs.get("bookmark_name")
        if bookmark_name:
            self._add_bookmark(para, bookmark_name)

    def _add_math_as_image(self, doc, element):
        """Render math formula as image using Playwright screenshot."""
        # Use counter to get the corresponding element ID
        selector = f"#math_element_{self.math_counter}"
        self.math_counter += 1

        # Generate cache key from content
        content = element.get_text().strip()
        content_hash = hashlib.md5(content.encode()).hexdigest()

        # Take screenshot
        screenshot_bytes = self._screenshot_element(selector, content_hash)

        if screenshot_bytes:
            # Add image to document
            image_stream = BytesIO(screenshot_bytes)
            try:
                doc.add_picture(image_stream, width=Inches(4))
                logger.debug("Added math formula as image: %s", selector)
            except Exception as e:
                logger.warning("Failed to add math image: %s", e)
                self._add_math_as_text(doc, element)
        else:
            # Fallback to text if screenshot fails
            self._add_math_as_text(doc, element)

    def _add_math_as_text(self, doc, element):
        """Fallback: Add math formula as formatted text."""
        formula = self._extract_math_text(element)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0

        # Block vs inline
        if element.name == "div":
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = p.add_run(formula)
        run.font.name = "Cambria Math"
        run.font.size = Pt(12)
        run.italic = True
        run.font.color.rgb = RGBColor(0, 0, 139)

    def _add_inline_math_image(self, paragraph, element):
        """Add inline math formula as an image within a paragraph run."""
        # Use counter to get the corresponding element ID
        selector = f"#math_element_{self.math_counter}"
        self.math_counter += 1

        # Generate cache key from content
        content = element.get_text().strip()
        content_hash = hashlib.md5(content.encode()).hexdigest()

        # Take screenshot
        screenshot_bytes = self._screenshot_element(selector, content_hash)

        if screenshot_bytes:
            # Add image inline in the paragraph
            try:
                run = paragraph.add_run()
                image_stream = BytesIO(screenshot_bytes)
                # Inline math at proper text height (12pt = 0.167 inches)
                # Slightly larger for readability: 0.18 inches
                run.add_picture(image_stream, height=Inches(0.18))
                logger.debug("Added inline math formula as image: %s", selector)
            except Exception as e:
                logger.warning("Failed to add inline math image: %s", e)
                # Fallback: add as text
                run = paragraph.add_run(content)
                run.font.name = "Cambria Math"
                run.italic = True
        else:
            # Fallback to text if screenshot fails
            run = paragraph.add_run(content)
            run.font.name = "Cambria Math"
            run.italic = True

    def _add_mermaid_as_image(self, doc, element):
        """Render Mermaid diagram as image using Playwright screenshot."""
        # Use counter to get the corresponding element ID
        selector = f"#mermaid_element_{self.mermaid_counter}"
        self.mermaid_counter += 1

        # Generate cache key from content
        content = element.get_text().strip()
        content_hash = hashlib.md5(content.encode()).hexdigest()

        # Take screenshot
        screenshot_bytes = self._screenshot_element(selector, content_hash)

        if screenshot_bytes:
            image_stream = BytesIO(screenshot_bytes)
            try:
                doc.add_picture(image_stream, width=Inches(6))
                logger.debug("Added Mermaid diagram as image: %s", selector)
            except Exception as e:
                logger.warning("Failed to add mermaid image: %s", e)
                self._add_diagram_as_code(doc, element)
        else:
            # Fallback to code display
            self._add_diagram_as_code(doc, element)

    def _add_diagram_as_code(self, doc, element):
        """Fallback: Add diagram code in formatted box."""
        code = element.get_text().strip()

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.add_run("[Mermaid Diagram]\n").bold = True
        p.add_run(f"Code:\n{code}")
        p.style = "No Spacing"

        self._add_shading(p, "FFFACD")
        self._add_paragraph_border(p)

    def _add_paragraph_with_formatting(self, doc, element, style=None):
        """Add paragraph with inline formatting."""
        logger.info(
            ">>> PARAGRAPH: has_content=%s, children_count=%d",
            self._has_content(element),
            len(list(element.children)) if hasattr(element, "children") else 0,
        )
        if not self._has_content(element):
            logger.info(">>> SKIPPING paragraph (no content)")
            return

        p = doc.add_paragraph(style=style)
        # Set tight spacing to match browser rendering
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        logger.info(">>> CALLING _add_formatted_content...")
        self._add_formatted_content(p, element)

    def _add_formatted_content(self, paragraph, element):
        """Recursively add inline formatted content."""
        logger.info(
            "_add_formatted_content called for element: %s",
            element.name if hasattr(element, "name") else type(element),
        )
        for child in element.children:
            logger.info(
                "  Processing child: %s", child.name if hasattr(child, "name") else type(child)
            )
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    paragraph.add_run(text)
            # Check for inline math (KaTeX rendered elements)
            elif child.name == "span" and self._has_math_class(child):
                # Screenshot inline math and add as inline image
                self._add_inline_math_image(paragraph, child)
            elif child.name in ["strong", "b"]:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ["em", "i"]:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "code":
                # Inline code
                run = paragraph.add_run(child.get_text())
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(214, 51, 132)
            elif child.name == "a":
                self._add_hyperlink(paragraph, child.get("href", ""), child.get_text())
            elif child.name == "br":
                paragraph.add_run("\n")
            elif child.name == "img":
                # Check if this is an emoji image - skip it and use alt text instead
                img_class = child.get("class", [])
                if isinstance(img_class, list) and "gemoji" in img_class:
                    # Emoji image - just use the alt text (Unicode emoji)
                    alt_text = child.get("alt", "")
                    if alt_text:
                        paragraph.add_run(alt_text)
                    continue

                # Add inline image (not just alt text)
                src = child.get("src", "")
                alt_text = child.get("alt", "")
                logger.info(
                    ">>> INLINE IMAGE DETECTED: alt='%s', src_length=%d",
                    alt_text,
                    len(src) if src else 0,
                )
                logger.info(">>> src prefix: %s", src[:100] if src else "EMPTY")
                if src:
                    image_added = False
                    try:
                        # Handle different image sources
                        if src.startswith("data:image"):
                            logger.info(">>> Extracting data URL image: %s", alt_text)
                            # Data URL - extract and add
                            image_bytes = self._extract_image_from_data_url(src)
                            if image_bytes:
                                logger.info(">>> Successfully extracted %d bytes", len(image_bytes))
                                try:
                                    run = paragraph.add_run()
                                    logger.info(">>> Created run, adding picture...")
                                    run.add_picture(BytesIO(image_bytes), width=Inches(4))
                                    image_added = True
                                    logger.info(
                                        ">>> ✅ Image added successfully to Word: %s", alt_text
                                    )
                                except Exception as pic_err:
                                    logger.error(
                                        ">>> ❌ Failed to add picture to run: %s",
                                        pic_err,
                                        exc_info=True,
                                    )
                            else:
                                logger.error(
                                    ">>> ❌ _extract_image_from_data_url returned None/empty"
                                )
                        elif src.startswith(("http://", "https://")):
                            logger.info("Downloading remote image: %s", src)
                            # Remote image - download and add
                            with urllib.request.urlopen(src) as response:
                                image_bytes = response.read()
                                run = paragraph.add_run()
                                run.add_picture(BytesIO(image_bytes), width=Inches(4))
                                image_added = True
                        elif src.startswith("/api/image?"):
                            # Local image served via API - extract file path
                            parsed = urllib.parse.urlparse(src)
                            params = urllib.parse.parse_qs(parsed.query)
                            if "path" in params:
                                file_path = params["path"][0]
                                logger.info("Loading /api/image file: %s", file_path)
                                if os.path.exists(file_path):
                                    run = paragraph.add_run()
                                    run.add_picture(file_path, width=Inches(4))
                                    image_added = True
                        else:
                            # Direct file path
                            file_path = src.replace("file:///", "").replace("/", os.sep)
                            logger.info("Loading direct file path: %s", file_path)
                            if os.path.exists(file_path):
                                run = paragraph.add_run()
                                run.add_picture(file_path, width=Inches(4))
                                image_added = True

                        if not image_added:
                            logger.warning(
                                "Image was not added (no matching handler or file not found): %s",
                                alt_text,
                            )
                            paragraph.add_run(f"[Image: {alt_text or 'missing'}]")
                    except Exception as e:
                        logger.error(
                            "Failed to add inline image '%s': %s", alt_text, e, exc_info=True
                        )
                        # Fallback to alt text - ALWAYS add something
                        paragraph.add_run(f"[Image: {alt_text or 'error'}]")
            else:
                if hasattr(child, "children"):
                    self._add_formatted_content(paragraph, child)

    def _add_list_item(self, doc, li_element, style):
        """Add list item with formatting."""
        p = doc.add_paragraph(style=style)
        # Set tight spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        self._add_formatted_content(p, li_element)

    def _add_code_block(self, doc, pre_element):
        """Add code block with styling."""
        code_element = pre_element.find("code")
        code_text = code_element.get_text() if code_element else pre_element.get_text()

        p = doc.add_paragraph()
        p.style = "No Spacing"
        # Set tight spacing
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

        self._add_shading(p, "F5F5F5")
        self._add_paragraph_border(p)

    def _add_image(self, doc, img_element):
        """Add image from various sources."""
        # Check if this is an emoji image - skip it (emojis already rendered as Unicode text)
        img_class = img_element.get("class", [])
        if isinstance(img_class, list) and "gemoji" in img_class:
            logger.debug("Skipping emoji image (already rendered as Unicode)")
            return

        src = img_element.get("src", "")
        alt_text = img_element.get("alt", "")

        logger.info("=== _add_image called: alt='%s', src='%s'", alt_text, src[:100] if src else "")

        if not src:
            if alt_text:
                p = doc.add_paragraph(f"[Image: {alt_text}]")
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.0
            logger.warning("Image has no src: %s", alt_text)
            return

        try:
            if src.startswith("data:image"):
                logger.info("Block image is data URL, extracting: %s", alt_text)
                self._add_image_from_data_url(doc, src)
                logger.info("Successfully added data URL image: %s", alt_text)
            elif src.startswith("/api/image?"):
                # Local image served via API - extract file path from query params
                logger.info("Block image is /api/image URL: %s", alt_text)
                parsed = urllib.parse.urlparse(src)
                params = urllib.parse.parse_qs(parsed.query)
                if "path" in params:
                    file_path = params["path"][0]
                    logger.info("Extracted path: %s", file_path)
                    if os.path.exists(file_path):
                        doc.add_picture(file_path, width=Inches(6))
                        logger.info("Added /api/image file: %s", file_path)
                    else:
                        logger.warning("File not found: %s", file_path)
                        p = doc.add_paragraph(f"[Image not found: {file_path}]")
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.0
            elif src.startswith(("http://", "https://")):
                # Remote image - download and embed
                logger.info("Downloading remote image: %s", src)
                with urllib.request.urlopen(src) as response:
                    image_bytes = response.read()
                    doc.add_picture(BytesIO(image_bytes), width=Inches(6))
                logger.info("Added remote image: %s", alt_text)
            else:
                # Direct file path
                file_path = src.replace("file:///", "").replace("/", os.sep)
                logger.info("Block image is direct file path: %s", file_path)
                if os.path.exists(file_path):
                    doc.add_picture(file_path, width=Inches(6))
                    logger.info("Added direct file image: %s", file_path)
                else:
                    logger.warning("Direct file not found: %s", file_path)
                    p = doc.add_paragraph(f"[Image: {alt_text or file_path}]")
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.0
        except Exception as e:
            logger.error("Could not load image '%s': %s", alt_text, e, exc_info=True)
            p = doc.add_paragraph(f"[Image: {alt_text or src}]")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0

    def _add_image_from_data_url(self, doc, data_url):
        """Extract and add image from data URL."""
        try:
            logger.info("Decoding data URL, length=%d", len(data_url))
            match = re.match(r"data:image/([^;]+);base64,(.+)", data_url)
            if match:
                image_format = match.group(1)
                base64_data = match.group(2)
                logger.info("Matched format=%s, base64_length=%d", image_format, len(base64_data))
                image_data = base64.b64decode(base64_data)
                logger.info("Decoded %d bytes of image data", len(image_data))
                image_stream = BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(6))
                logger.info("Successfully added picture to doc")
            else:
                logger.warning("Data URL pattern did not match: %s", data_url[:100])
        except Exception as e:
            logger.error("Could not decode data URL: %s", e, exc_info=True)

    def _extract_image_from_data_url(self, data_url):
        """Extract image bytes from data URL."""
        try:
            match = re.match(r"data:image/([^;]+);base64,(.+)", data_url)
            if match:
                image_format = match.group(1)
                base64_data = match.group(2)
                logger.info(
                    "Decoding base64 image, format=%s, data_length=%d",
                    image_format,
                    len(base64_data),
                )
                image_bytes = base64.b64decode(base64_data)
                logger.info("Successfully decoded %d bytes", len(image_bytes))
                return image_bytes
            else:
                logger.warning("Data URL does not match expected pattern: %s", data_url[:100])
        except Exception as e:
            logger.error("Could not decode data URL: %s", e, exc_info=True)
        return None

    def _add_horizontal_rule(self, doc):
        """Add horizontal line separator."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("_" * 80)
        run.font.color.rgb = RGBColor(192, 192, 192)

    def _add_hyperlink(self, paragraph, url, text):
        """Add clickable external hyperlink."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)  # pylint: disable=protected-access

    def _add_shading(self, paragraph, color_hex):
        """Add background color to paragraph."""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color_hex)
        paragraph._p.get_or_add_pPr().append(shading_elm)  # pylint: disable=protected-access

    def _add_paragraph_border(self, paragraph):
        """Add border around paragraph."""
        pPr = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
        pBdr = OxmlElement("w:pBdr")

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), "CCCCCC")
            pBdr.append(border)

        pPr.append(pBdr)

    def _process_table(self, doc, table_element):
        """Process HTML table."""
        rows = table_element.find_all("tr")
        if not rows:
            return

        max_cols = max(len(row.find_all(["th", "td"])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Light Grid Accent 1"

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for j, cell in enumerate(cells):
                cell_para = table.rows[i].cells[j].paragraphs[0]
                cell_para.clear()
                self._add_formatted_content(cell_para, cell)

                if cell.name == "th":
                    for run in cell_para.runs:
                        run.font.bold = True

    def _get_text_with_emojis(self, element):
        """Extract text with emoji preservation."""
        return element.get_text()

    def _has_content(self, element):
        """Check if element has content."""
        # Check for text content
        if element.get_text().strip():
            return True
        # Check for img tags (even if no text)
        if element.find("img"):
            return True
        return False

    def _has_math_class(self, element):
        """Check for math-related classes."""
        classes = element.get("class", [])
        return any(
            cls in classes for cls in ["katex", "katex-display", "math", "mathjax", "arithmatex"]
        )

    def _extract_math_text(self, element):
        """Extract LaTeX from math element (fallback)."""
        # Try annotation first
        annotation = element.find("annotation", {"encoding": "application/x-tex"})
        if annotation:
            return annotation.get_text().strip()

        # Try math element
        math_element = element.find("math")
        if math_element:
            annotation = math_element.find("annotation")
            if annotation:
                return annotation.get_text().strip()

        # Clean text content
        text = element.get_text().strip()
        if not text:
            return "[Math Formula]"

        # Remove delimiters
        text = re.sub(r"^\\\[(.*)\\\]$", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"^\\\((.*)\\\)$", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"^\$\$(.*)\$\$$", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"^\$(.*)\$$", r"\1", text, flags=re.DOTALL)
        text = " ".join(text.split())

        return text if text else "[Math Formula]"
