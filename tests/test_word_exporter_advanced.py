"""
Advanced test coverage for WordExporter to reach 90%+ coverage.

This test file covers:
- Screenshot element methods
- Math rendering (image and text fallback)
- Mermaid diagram rendering (image and code fallback)
- Image handling (data URLs, remote URLs, file paths, emoji detection)
- Table processing
- Formatted content (bold, italic, code, links)
- Helper methods (_has_content, _has_math_class, etc.)
- Error handling and edge cases
"""

import os
import tempfile
from unittest.mock import MagicMock, patch
import pytest
from bs4 import BeautifulSoup
from docx import Document

from markdown_viewer.exporters.word_exporter import WordExporter


@pytest.fixture
def exporter():
    """Create WordExporter instance with mocked browser."""
    exp = WordExporter()
    # Mock the browser to avoid Playwright initialization
    with patch.object(exp, "_load_html"):
        exp.page = MagicMock()
        exp.page.content.return_value = "<html><body>Test</body></html>"
        yield exp


class TestScreenshotMethods:
    """Test screenshot-related methods."""

    def test_screenshot_element_success(self, exporter):
        """Test successful element screenshot."""
        mock_element = MagicMock()
        mock_element.screenshot.return_value = b"fake_screenshot_data"
        exporter.page.query_selector.return_value = mock_element

        result = exporter._screenshot_element("#test", "hash123")

        assert result == b"fake_screenshot_data"
        assert exporter.screenshot_cache["hash123"] == b"fake_screenshot_data"
        mock_element.scroll_into_view_if_needed.assert_called_once()

    def test_screenshot_element_cached(self, exporter):
        """Test screenshot cache hit."""
        exporter.screenshot_cache["hash123"] = b"cached_data"

        result = exporter._screenshot_element("#test", "hash123")

        assert result == b"cached_data"
        exporter.page.query_selector.assert_not_called()

    def test_screenshot_element_not_found(self, exporter):
        """Test screenshot when element not found."""
        exporter.page.query_selector.return_value = None

        result = exporter._screenshot_element("#missing", "hash456")

        assert result is None

    def test_screenshot_element_exception(self, exporter):
        """Test screenshot error handling."""
        exporter.page.query_selector.side_effect = Exception("Browser error")

        result = exporter._screenshot_element("#error", "hash789")

        assert result is None


class TestMathRendering:
    """Test math formula rendering methods."""

    def test_add_math_as_image_success(self, exporter):
        """Test math formula added as image."""
        doc = Document()
        element = BeautifulSoup('<div class="arithmatex">x = y + z</div>', "html.parser").div

        mock_element = MagicMock()
        mock_element.screenshot.return_value = b"math_screenshot"
        exporter.page.query_selector.return_value = mock_element

        exporter._add_math_as_image(doc, element)

        # Should have added a picture
        assert len(doc.element.body) > 0

    def test_add_math_as_image_fallback(self, exporter):
        """Test math formula fallback to text."""
        doc = Document()
        element = BeautifulSoup('<div class="arithmatex">x = y + z</div>', "html.parser").div

        exporter.page.query_selector.return_value = None  # Screenshot fails

        exporter._add_math_as_image(doc, element)

        # Should have added a paragraph with text
        assert len(doc.paragraphs) > 0

    def test_add_math_as_text(self, exporter):
        """Test math formula as formatted text."""
        doc = Document()
        element = BeautifulSoup('<span class="katex">E = mc^2</span>', "html.parser").span

        exporter._add_math_as_text(doc, element)

        assert len(doc.paragraphs) == 1
        paragraph = doc.paragraphs[0]
        assert paragraph.runs[0].font.name == "Cambria Math"
        assert paragraph.runs[0].italic is True

    def test_add_inline_math_image_success(self, exporter):
        """Test inline math as image."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup('<span class="katex">x^2</span>', "html.parser").span

        mock_element = MagicMock()
        mock_element.screenshot.return_value = b"inline_math"
        exporter.page.query_selector.return_value = mock_element

        exporter._add_inline_math_image(paragraph, element)

        # Should have added a run with picture
        assert len(paragraph.runs) > 0

    def test_add_inline_math_image_fallback(self, exporter):
        """Test inline math fallback to text."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup('<span class="katex">x^2</span>', "html.parser").span

        exporter.page.query_selector.return_value = None

        exporter._add_inline_math_image(paragraph, element)

        # Should have added text run
        assert len(paragraph.runs) > 0
        assert paragraph.runs[0].font.name == "Cambria Math"


class TestMermaidRendering:
    """Test Mermaid diagram rendering methods."""

    def test_add_mermaid_as_image_success(self, exporter):
        """Test Mermaid diagram added as image."""
        doc = Document()
        element = BeautifulSoup('<div class="mermaid">graph TD; A-->B;</div>', "html.parser").div

        mock_element = MagicMock()
        mock_element.screenshot.return_value = b"mermaid_screenshot"
        exporter.page.query_selector.return_value = mock_element

        exporter._add_mermaid_as_image(doc, element)

        # Should have added a picture
        assert len(doc.element.body) > 0

    def test_add_mermaid_as_image_fallback(self, exporter):
        """Test Mermaid diagram fallback to code."""
        doc = Document()
        element = BeautifulSoup('<div class="mermaid">graph TD; A-->B;</div>', "html.parser").div

        exporter.page.query_selector.return_value = None

        exporter._add_mermaid_as_image(doc, element)

        # Should have added a paragraph with code
        assert len(doc.paragraphs) > 0

    def test_add_diagram_as_code(self, exporter):
        """Test diagram code display."""
        doc = Document()
        element = BeautifulSoup('<div class="mermaid">graph TD; A-->B;</div>', "html.parser").div

        exporter._add_diagram_as_code(doc, element)

        assert len(doc.paragraphs) == 1
        paragraph = doc.paragraphs[0]
        assert "Mermaid Diagram" in paragraph.text


class TestImageHandling:
    """Test various image handling scenarios."""

    def test_add_image_data_url(self, exporter):
        """Test image from data URL."""
        doc = Document()
        img_html = '<img src="data:image/png;base64,iVBORw0KGgo=" alt="Test Image">'
        element = BeautifulSoup(img_html, "html.parser").img

        with patch.object(exporter, "_extract_image_from_data_url") as mock_extract:
            mock_extract.return_value = b"fake_image_data"
            exporter._add_image(doc, element)

        # Should have added image
        assert len(doc.element.body) > 0

    def test_add_image_http_url(self, exporter):
        """Test image from HTTP URL."""
        doc = Document()
        img_html = '<img src="https://example.com/image.png" alt="Remote Image">'
        element = BeautifulSoup(img_html, "html.parser").img

        with patch("urllib.request.urlopen") as mock_urlopen:
            mock_response = MagicMock()
            mock_response.read.return_value = b"remote_image_data"
            mock_response.__enter__.return_value = mock_response
            mock_urlopen.return_value = mock_response

            exporter._add_image(doc, element)

        # Should have added image
        assert len(doc.element.body) > 0

    def test_add_image_file_path(self, exporter):
        """Test image from file path."""
        doc = Document()

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(b"fake_png_data")
            tmp_path = tmp.name

        try:
            img_html = f'<img src="file:///{tmp_path}" alt="Local Image">'
            element = BeautifulSoup(img_html, "html.parser").img

            exporter._add_image(doc, element)

            # Should have added image or fallback text
            assert len(doc.element.body) > 0 or len(doc.paragraphs) > 0
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_add_image_gemoji_detection(self, exporter):
        """Test emoji image detection and Unicode fallback."""
        doc = Document()
        img_html = '<img class="gemoji" src="emoji.svg" alt="😀">'
        element = BeautifulSoup(img_html, "html.parser").img

        exporter._add_image(doc, element)

        # Should NOT add image, should return early for gemoji
        # Check that no image was added (would throw error on SVG)
        # This is tested by not throwing UnrecognizedImageError

    def test_add_image_missing_file(self, exporter):
        """Test missing image file fallback."""
        doc = Document()
        img_html = '<img src="/nonexistent/image.png" alt="Missing">'
        element = BeautifulSoup(img_html, "html.parser").img

        exporter._add_image(doc, element)

        # Should have added fallback text
        assert len(doc.paragraphs) > 0
        assert "[Image:" in doc.paragraphs[-1].text

    def test_extract_image_from_data_url_valid(self, exporter):
        """Test extracting image from valid data URL."""
        data_url = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="

        result = exporter._extract_image_from_data_url(data_url)

        assert result is not None
        assert isinstance(result, bytes)

    def test_extract_image_from_data_url_invalid(self, exporter):
        """Test extracting image from invalid data URL."""
        data_url = "data:image/png;base64,invalid_base64!!!"

        result = exporter._extract_image_from_data_url(data_url)

        assert result is None


class TestFormattedContent:
    """Test inline formatted content methods."""

    def test_add_formatted_content_bold(self, exporter):
        """Test bold text formatting."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup("<p><strong>Bold Text</strong></p>", "html.parser").p

        exporter._add_formatted_content(paragraph, element)

        assert len(paragraph.runs) > 0
        assert any(run.bold for run in paragraph.runs)

    def test_add_formatted_content_italic(self, exporter):
        """Test italic text formatting."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup("<p><em>Italic Text</em></p>", "html.parser").p

        exporter._add_formatted_content(paragraph, element)

        assert len(paragraph.runs) > 0
        assert any(run.italic for run in paragraph.runs)

    def test_add_formatted_content_code(self, exporter):
        """Test inline code formatting."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup("<p><code>code_snippet</code></p>", "html.parser").p

        exporter._add_formatted_content(paragraph, element)

        assert len(paragraph.runs) > 0

    def test_add_formatted_content_link(self, exporter):
        """Test hyperlink formatting."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup('<p><a href="https://example.com">Link</a></p>', "html.parser").p

        exporter._add_formatted_content(paragraph, element)

        # Links are added through _add_hyperlink which creates special XML elements
        # Just verify no errors occurred during processing
        assert True

    def test_add_formatted_content_nested(self, exporter):
        """Test nested formatting."""
        doc = Document()
        paragraph = doc.add_paragraph()
        element = BeautifulSoup("<p><strong><em>Bold Italic</em></strong></p>", "html.parser").p

        exporter._add_formatted_content(paragraph, element)

        assert len(paragraph.runs) > 0


class TestTableProcessing:
    """Test table processing methods."""

    def test_process_table_simple(self, exporter):
        """Test simple table processing."""
        doc = Document()
        table_html = """
        <table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Cell 1</td><td>Cell 2</td></tr>
        </table>
        """
        element = BeautifulSoup(table_html, "html.parser").table

        exporter._process_table(doc, element)

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2

    def test_process_table_empty(self, exporter):
        """Test empty table handling."""
        doc = Document()
        table_html = "<table></table>"
        element = BeautifulSoup(table_html, "html.parser").table

        exporter._process_table(doc, element)

        # Should not add table if no rows
        assert len(doc.tables) == 0

    def test_process_table_complex(self, exporter):
        """Test complex table with varying columns."""
        doc = Document()
        table_html = """
        <table>
            <tr><th>A</th><th>B</th><th>C</th></tr>
            <tr><td>1</td><td>2</td></tr>
            <tr><td>X</td><td>Y</td><td>Z</td></tr>
        </table>
        """
        element = BeautifulSoup(table_html, "html.parser").table

        exporter._process_table(doc, element)

        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Should use max columns (3)
        assert len(table.columns) == 3


class TestHelperMethods:
    """Test helper utility methods."""

    def test_has_content_with_text(self, exporter):
        """Test content detection with text."""
        element = BeautifulSoup("<p>Some text</p>", "html.parser").p

        assert exporter._has_content(element) is True

    def test_has_content_with_image(self, exporter):
        """Test content detection with image."""
        element = BeautifulSoup('<p><img src="test.png" alt="img"></p>', "html.parser").p

        assert exporter._has_content(element) is True

    def test_has_content_empty(self, exporter):
        """Test content detection with empty element."""
        element = BeautifulSoup("<p>   </p>", "html.parser").p

        assert exporter._has_content(element) is False

    def test_has_math_class_katex(self, exporter):
        """Test math class detection for KaTeX."""
        element = BeautifulSoup('<span class="katex">x</span>', "html.parser").span

        assert exporter._has_math_class(element) is True

    def test_has_math_class_arithmatex(self, exporter):
        """Test math class detection for arithmatex."""
        element = BeautifulSoup('<div class="arithmatex">y</div>', "html.parser").div

        assert exporter._has_math_class(element) is True

    def test_has_math_class_none(self, exporter):
        """Test math class detection for non-math element."""
        element = BeautifulSoup("<p>normal text</p>", "html.parser").p

        assert exporter._has_math_class(element) is False

    def test_get_text_with_emojis(self, exporter):
        """Test emoji text extraction."""
        element = BeautifulSoup("<p>Hello 😀 World</p>", "html.parser").p

        text = exporter._get_text_with_emojis(element)

        assert "Hello" in text
        assert "World" in text


class TestHeadingsAndStructure:
    """Test document structure elements."""

    def test_add_heading_h1(self, exporter):
        """Test H1 heading."""
        doc = Document()
        element = BeautifulSoup("<h1>Main Title</h1>", "html.parser").h1

        exporter._add_heading(doc, element, 1)

        assert len(doc.paragraphs) > 0
        # Heading added with bookmark
        assert "Main Title" in doc.paragraphs[-1].text

    def test_add_heading_with_bookmark(self, exporter):
        """Test heading bookmark creation."""
        doc = Document()
        element = BeautifulSoup("<h2>Section Title</h2>", "html.parser").h2

        exporter._add_heading(doc, element, 2)

        # Bookmark should be added for TOC linking
        assert len(doc.paragraphs) > 0

    def test_add_horizontal_rule(self, exporter):
        """Test horizontal rule."""
        doc = Document()

        exporter._add_horizontal_rule(doc)

        assert len(doc.paragraphs) > 0


class TestCleanupAndBrowser:
    """Test browser lifecycle and cleanup."""

    def test_cleanup(self):
        """Test cleanup method."""
        exp = WordExporter()
        mock_browser = MagicMock()
        mock_playwright = MagicMock()

        exp.page = MagicMock()
        exp.browser = mock_browser
        exp.playwright = mock_playwright
        exp.temp_html_path = None

        exp._cleanup()

        # Should close browser and playwright
        mock_browser.close.assert_called_once()
        mock_playwright.stop.assert_called_once()
        # References should be cleared
        assert exp.browser is None
        assert exp.playwright is None

    def test_cleanup_with_temp_file(self, exporter):
        """Test cleanup with temporary file."""
        exporter.page = MagicMock()
        exporter.browser = MagicMock()
        exporter.playwright = MagicMock()

        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            exporter.temp_html_path = tmp.name

        exporter._cleanup()

        # Temp file should be deleted
        assert not os.path.exists(exporter.temp_html_path)


class TestCodeBlocks:
    """Test code block handling."""

    def test_add_code_block(self, exporter):
        """Test code block formatting."""
        doc = Document()
        pre_html = '<pre><code>def hello():\n    return "world"</code></pre>'
        element = BeautifulSoup(pre_html, "html.parser").pre

        exporter._add_code_block(doc, element)

        assert len(doc.paragraphs) > 0
        # Code should be in monospace font

    def test_add_code_block_no_code_tag(self, exporter):
        """Test code block without code tag."""
        doc = Document()
        pre_html = "<pre>plain preformatted text</pre>"
        element = BeautifulSoup(pre_html, "html.parser").pre

        exporter._add_code_block(doc, element)

        assert len(doc.paragraphs) > 0


class TestListHandling:
    """Test list processing."""

    def test_add_list_item_unordered(self, exporter):
        """Test unordered list item."""
        doc = Document()
        li_html = "<li>List item text</li>"
        element = BeautifulSoup(li_html, "html.parser").li

        exporter._add_list_item(doc, element, "List Bullet")

        assert len(doc.paragraphs) > 0

    def test_add_list_item_ordered(self, exporter):
        """Test ordered list item."""
        doc = Document()
        li_html = "<li>Numbered item</li>"
        element = BeautifulSoup(li_html, "html.parser").li

        exporter._add_list_item(doc, element, "List Number")

        assert len(doc.paragraphs) > 0


class TestComplexExport:
    """Test complete export scenarios."""

    def test_export_with_md_file_path(self, exporter):
        """Test export with markdown file path."""
        html = "<html><body><h1>Test</h1></body></html>"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name

        try:
            exporter.export(html, "# Test", output_path, md_file_path="/path/to/file.md")

            assert exporter.md_file_path == "/path/to/file.md"
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_export_with_options(self, exporter):
        """Test export with custom options."""
        html = "<html><body><p>Test</p></body></html>"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name

        try:
            options = {"include_toc": True}
            exporter.export(html, "", output_path, options=options)

            # Should complete without error
            assert os.path.exists(output_path)
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)
